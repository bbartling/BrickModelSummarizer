import os
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from pydantic import BaseModel
from openai import OpenAI
from docx import Document
import json
import docx

# Ensure the OpenAI API key is set
api_key = os.getenv("OPENAI_API_KEY")
if not api_key:
    raise EnvironmentError("OPENAI_API_KEY environment variable is not set.")

# Initialize the OpenAI client
client = OpenAI(api_key=api_key)

# Define structured response schemas
class Step(BaseModel):
    explanation: str
    output: str

class FaultDiagnostics(BaseModel):
    steps: list[Step]
    final_analysis: str

# Fault details
fault_equation = (
    "fault_flag = 1 if (MAT < min(RAT, OAT) or MAT > max(RAT, OAT)) and (VFDSPD > 0.01) else 0"
)
fault_description = (
    "Fault Condition: Mix air temperature (MAT) must be between return air temperature (RAT) "
    "and outside air temperature (OAT) when the fan is running."
)
required_inputs = (
    "- Mix air temperature (MAT)\n"
    "- Return air temperature (RAT)\n"
    "- Outside air temperature (OAT)\n"
    "- Supply fan VFD speed (VFDSPD)"
)

# Preprocess and apply fault detection
def detect_faults(df, mat_col, rat_col, oat_col, fan_speed_col):
    df["fault_flag"] = (
        (df[mat_col] < df[[rat_col, oat_col]].min(axis=1)) |
        (df[mat_col] > df[[rat_col, oat_col]].max(axis=1))
    ) & (df[fan_speed_col] > 0.01)
    return df

# Generate summary statistics
def generate_summary(df, fault_col, fan_speed_col, outlier_mask):
    return {
        "total_data_points": int(len(df)),
        "fan_running_points": int(df[fan_speed_col].gt(0.01).sum()),
        "fault_count": int(df[fault_col].sum()),
        "outlier_count": int(outlier_mask.any(axis=1).sum()),
        "sensor_stats": df.describe().to_dict(),
    }

# Detect outliers
def detect_outliers(df, columns):
    outlier_mask = pd.DataFrame(False, index=df.index, columns=columns)
    for col in columns:
        outlier_mask[col] = df[col] <= 0
    return outlier_mask

# Clean data by removing outliers
def clean_outliers(df, columns):
    return df[~((df[columns] <= 0).any(axis=1))]

# Plot faults
def plot_faults(df, fault_col, output_path):
    plt.figure(figsize=(10, 6))
    df[fault_col].cumsum().plot(title="Cumulative Faults Over Time", xlabel="Timestamp", ylabel="Cumulative Faults")
    plt.grid()
    plt.savefig(output_path)
    plt.close()
    print(f"Plot saved to {output_path}")

# Box plot of temperature ranges
def plot_temperature_ranges(df, mat_col, rat_col, oat_col, output_path):
    plt.figure(figsize=(10, 6))
    df[[mat_col, rat_col, oat_col]].plot(kind='box', title="Temperature Ranges (MAT, RAT, OAT)")
    plt.ylabel("Temperature (°F)")
    plt.savefig(output_path)
    plt.close()
    print(f"Temperature ranges plot saved to {output_path}")

# Histogram of hourly faults
def create_hourly_fault_histogram(df, fault_col, output_path):
    df[f"hour_of_the_day_{fault_col}"] = df.index.hour.where(df[fault_col] == 1)
    plt.figure(figsize=(12, 6))
    plt.hist(df[f"hour_of_the_day_{fault_col}"].dropna(), bins=range(0, 25), edgecolor='black', align='left')
    plt.xticks(range(0, 24))
    plt.xlabel("Hour of the Day")
    plt.ylabel("Frequency")
    plt.title(f"Hour-of-Day When Fault Flag {fault_col} is TRUE")
    plt.savefig(output_path)
    plt.close()
    print(f"Hourly fault histogram saved to {output_path}")

# Log LLM decision-making steps
def log_llm_steps(doc, llm_steps):
    doc.add_heading("LLM Decision-Making Steps", level=2)
    for step in llm_steps:
        doc.add_paragraph(step)

# Save report to Word document
def save_report_to_docx_with_histogram(response, summary, fault_plot_path, temp_plot_path, hist_plot_path, output_path):
    doc = Document()
    doc.add_heading("HVAC Diagnostics Report", level=1)
    doc.add_heading("Fault Rule Details", level=2)
    doc.add_paragraph(f"Equation:\n{fault_equation}")
    doc.add_paragraph(f"Description:\n{fault_description}")
    doc.add_paragraph(f"Required Inputs:\n{required_inputs}")
    doc.add_heading("Summary Statistics", level=2)
    doc.add_paragraph(f"Total Data Points: {summary['total_data_points']}")
    doc.add_paragraph(f"Fan Running Points: {summary['fan_running_points']}")
    doc.add_paragraph(f"Fault Count: {summary['fault_count']}")
    doc.add_paragraph(f"Outlier Count: {summary['outlier_count']}")
    doc.add_heading("OpenAI Diagnostics", level=2)
    if response:
        for step in response.steps:
            doc.add_paragraph(f"Explanation: {step.explanation}")
            doc.add_paragraph(f"Output: {step.output}")
        doc.add_paragraph(f"Final Analysis: {response.final_analysis}")
    else:
        doc.add_paragraph("No response from diagnostics.")
    doc.add_heading("Visualizations", level=2)
    doc.add_paragraph("Cumulative Faults Over Time:")
    doc.add_picture(fault_plot_path, width=docx.shared.Inches(6))
    doc.add_paragraph("Temperature Ranges (MAT, RAT, OAT):")
    doc.add_picture(temp_plot_path, width=docx.shared.Inches(6))
    doc.add_paragraph("Hourly Fault Occurrences:")
    doc.add_picture(hist_plot_path, width=docx.shared.Inches(6))
    log_llm_steps(doc, llm_steps)
    doc.save(output_path)
    print(f"Report saved to {output_path}")

def query_openai_for_diagnostics(rule_description, summary):
    """Request structured diagnostics from OpenAI."""
    try:
        completion = client.beta.chat.completions.parse(
            model="gpt-4o-2024-08-06",
            messages=[
                {"role": "system", "content": "You are an HVAC fault detection and diagnostics expert."},
                {"role": "user", "content": (
                    f"As an HVAC FDD expert, analyze the following fault rule and dataset summary:\n\n"
                    f"Fault Rule: '{rule_description}'\n\n"
                    f"Dataset Summary:\n{json.dumps(summary, indent=2)}\n\n"
                    "Provide a structured analysis, including potential patterns, anomalies, "
                    "or false positives in the data."
                )}
            ],
            response_format=FaultDiagnostics,
        )
        return completion.choices[0].message.parsed
    except Exception as e:
        print(f"Error during OpenAI diagnostics query: {e}")
        return None

# Query OpenAI for triage
def query_openai_for_triage(rule_description, summary, fault_rate, outlier_count):
    class TriageDecision(BaseModel):
        action: str
        explanation: str

    try:
        completion = client.beta.chat.completions.parse(
            model="gpt-4o-2024-08-06",
            messages=[
                {"role": "system", "content": "You are an HVAC diagnostics expert."},
                {"role": "user", "content": (
                    f"Fault Rule: '{rule_description}'\n\n"
                    f"Dataset Summary:\n{json.dumps(summary, indent=2)}\n\n"
                    f"Fault Rate: {fault_rate:.2%}\n\n"
                    f"Outlier Count: {outlier_count}\n\n"
                    "Based on this information, recommend an action."
                )}
            ],
            response_format=TriageDecision,
        )
        return completion.choices[0].message.parsed
    except Exception as e:
        print(f"Error during OpenAI triage query: {e}")
        return {"action": "Error", "explanation": str(e)}

# Main script
if __name__ == "__main__":
    filepath = r"C:\Users\bbartling\Downloads\Midtown_PD_Master.csv"
    fault_plot_output_path = r"C:\Users\bbartling\Downloads\HVAC_Faults_Plot.png"
    temp_plot_output_path = r"C:\Users\bbartling\Downloads\Temperature_Ranges_Plot.png"
    hist_plot_output_path = r"C:\Users\bbartling\Downloads\Hourly_Fault_Histogram.png"
    report_output_path = r"C:\Users\bbartling\Downloads\HVAC_Diagnostics_Report.docx"
    mat_col, rat_col, oat_col, fan_speed_col = "MA_Temp", "RA_Temp", "OaTemp", "SA_FanSpeed"

    df = pd.read_csv(filepath, parse_dates=["timestamp"]).set_index("timestamp")
    outlier_mask = detect_outliers(df, [mat_col, rat_col])
    df = clean_outliers(df, [mat_col, rat_col])[df[fan_speed_col] > 0.01]
    df = detect_faults(df, mat_col, rat_col, oat_col, fan_speed_col)
    summary = generate_summary(df, "fault_flag", fan_speed_col, outlier_mask)
    fault_rate = summary["fault_count"] / summary["total_data_points"]

    triage_response = query_openai_for_triage(fault_description, summary, fault_rate, summary["outlier_count"])
    llm_steps = [
        f"Fault detection: {summary['fault_count']} faults identified.",
        f"Outlier detection: {summary['outlier_count']} outliers identified.",
        f"Triage decision: {triage_response.action}. Explanation: {triage_response.explanation}.",
    ]
    plot_faults(df, "fault_flag", fault_plot_output_path)
    plot_temperature_ranges(df, mat_col, rat_col, oat_col, temp_plot_output_path)
    create_hourly_fault_histogram(df, "fault_flag", hist_plot_output_path)
    diagnostics_response = query_openai_for_diagnostics(fault_description, summary)
    save_report_to_docx_with_histogram(
        diagnostics_response, summary, fault_plot_output_path, temp_plot_output_path, hist_plot_output_path, report_output_path
    )
